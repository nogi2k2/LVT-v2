import asyncio
import ast
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from contextlib import contextmanager, suppress
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Set, Tuple
from xml.etree import ElementTree as ET

from fastmcp.exceptions import ToolError
from PIL import Image
from tqdm import tqdm
from ultrarag.server import UltraRAG_MCP_Server


def _validate_path(user_path: str, allowed_base: Optional[str] = None) -> Path:
    """Validate and sanitize file path to prevent path traversal attacks.
    
    Args:
        user_path: User-provided file path
        allowed_base: Optional base directory to restrict paths to
        
    Returns:
        Resolved and validated Path object
        
    Raises:
        ValueError: If path traversal is detected or path is invalid
    """
    try:
        # Resolve the path to absolute
        safe_path = Path(user_path).resolve()
        
        # If allowed_base is provided, ensure path is within it
        if allowed_base:
            base_path = Path(allowed_base).resolve()
            try:
                # Check if safe_path is relative to base_path
                safe_path.relative_to(base_path)
            except ValueError:
                raise ValueError(
                    f"Path traversal detected: '{user_path}' is outside allowed directory '{allowed_base}'"
                )
        
        # Additional safety: check for suspicious patterns
        path_str = str(safe_path)
        if ".." in path_str or path_str.startswith("/etc/") or path_str.startswith("/proc/"):
            # Double check even after resolve
            if ".." in str(Path(user_path)):
                raise ValueError(f"Path traversal detected: '{user_path}' contains '..'")
        
        return safe_path
    except (OSError, ValueError) as e:
        if isinstance(e, ValueError):
            raise
        raise ValueError(f"Invalid path: {user_path}") from e

app = UltraRAG_MCP_Server("corpus")

_ANSI_ESCAPE_RE = re.compile(r"\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])")
_TQDM_PROGRESS_RE = re.compile(r"^(?P<name>[^:]{1,120}):\s*(?P<pct>\d{1,3})%")


def _iter_clean_subprocess_lines(raw_line: bytes) -> Iterable[str]:
    """Yield cleaned log lines from a subprocess stdout chunk."""
    text = raw_line.decode("utf-8", errors="replace")
    text = text.replace("\r", "\n")
    text = _ANSI_ESCAPE_RE.sub("", text)
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if len(line) > 320:
            line = line[:317] + "..."
        yield line


def _extract_progress_update(line: str) -> Optional[Tuple[str, int]]:
    """Extract progress updates from tqdm-like lines."""
    if "|" not in line:
        return None
    match = _TQDM_PROGRESS_RE.match(line)
    if not match:
        return None
    name = match.group("name").strip()
    try:
        pct = int(match.group("pct"))
    except ValueError:
        return None
    return name, pct


def _can_render_live_tqdm() -> bool:
    """Whether current process should render local tqdm progress bars."""
    flag = os.getenv("ULTRARAG_MINERU_LIVE_TQDM", "1").strip().lower()
    if flag in {"0", "false", "no", "off"}:
        return False
    try:
        return sys.stderr.isatty()
    except Exception:
        return False


@contextmanager
def suppress_stdout():
    """Context manager to suppress stdout output."""
    stdout_fd = sys.stdout.fileno()
    saved_stdout_fd = os.dup(stdout_fd)

    try:
        devnull = os.open(os.devnull, os.O_WRONLY)
        os.dup2(devnull, stdout_fd)
        os.close(devnull)
        yield
    finally:
        os.dup2(saved_stdout_fd, stdout_fd)
        os.close(saved_stdout_fd)


def _local_name(tag: str) -> str:
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _read_docx_text_zip(fp: str) -> Optional[str]:
    try:
        with zipfile.ZipFile(fp) as zf:
            if "word/document.xml" not in zf.namelist():
                return None
            xml_bytes = zf.read("word/document.xml")
    except zipfile.BadZipFile:
        return None
    except Exception as e:
        app.logger.warning(f"Docx zip read failed: {fp} | {e}")
        return None

    try:
        root = ET.fromstring(xml_bytes)
    except Exception as e:
        app.logger.warning(f"Docx xml parse failed: {fp} | {e}")
        return None

    paras: List[str] = []
    for p in root.iter():
        if _local_name(p.tag) != "p":
            continue
        buf: List[str] = []
        for node in p.iter():
            lname = _local_name(node.tag)
            if lname == "t":
                if node.text:
                    buf.append(node.text)
            elif lname == "tab":
                buf.append("\t")
            elif lname in ("br", "cr"):
                buf.append("\n")
        para_text = "".join(buf).strip()
        if para_text:
            paras.append(para_text)
    return "\n".join(paras)


def _read_docx_text(fp: str) -> Optional[str]:
    docx_import_error: Optional[Exception] = None
    try:
        from docx import Document
    except ImportError as e:
        docx_import_error = e
    else:
        try:
            doc = Document(fp)
            full_text = [para.text for para in doc.paragraphs]
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    full_text.append(" | ".join(row_text))
            return "\n".join(full_text)
        except Exception as e:
            app.logger.warning(f"Docx read failed (python-docx): {fp} | {e}")

    content = _read_docx_text_zip(fp)
    if content is not None:
        return content

    if docx_import_error is not None:
        err_msg = "python-docx not installed and docx zip parse failed. Please `pip install python-docx`."
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    return None


def _find_office_cmd() -> Optional[str]:
    return shutil.which("soffice") or shutil.which("libreoffice")


def _convert_to_docx_with_office(fp: str, out_dir: str, office_cmd: str) -> Optional[str]:
    cmd = [
        office_cmd,
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        out_dir,
        fp,
    ]
    try:
        subprocess.run(
            cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            check=True,
            timeout=60,
        )
    except Exception as e:
        app.logger.warning(f"Office convert failed: {fp} | {e}")
        return None

    expected = Path(out_dir) / f"{Path(fp).stem}.docx"
    if expected.exists():
        return str(expected)
    for p in Path(out_dir).glob("*.docx"):
        return str(p)
    return None


def _read_via_office_convert(fp: str) -> Optional[str]:
    office_cmd = _find_office_cmd()
    if not office_cmd:
        return None
    with tempfile.TemporaryDirectory(prefix="ultrarag_docx_") as tmpdir:
        out_path = _convert_to_docx_with_office(fp, tmpdir, office_cmd)
        if not out_path:
            return None
        return _read_docx_text(out_path)


def _save_jsonl(rows: Iterable[Dict[str, Any]], file_path: str) -> None:
    """Save rows to a JSONL file.

    Args:
        rows: Iterable of dictionaries to save
        file_path: Path to the output JSONL file
    """
    out_dir = Path(file_path).parent
    if out_dir and str(out_dir) != ".":
        os.makedirs(out_dir, exist_ok=True)

    with open(file_path, "w", encoding="utf-8", newline="\n") as f:
        for r in rows:
            f.write(json.dumps(r, ensure_ascii=False) + "\n")


def _load_jsonl(file_path: str) -> List[Dict[str, Any]]:
    """Load documents from a JSONL file.

    Args:
        file_path: Path to the JSONL file

    Returns:
        List of dictionaries loaded from the file
    """
    docs = []
    with open(file_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            docs.append(json.loads(line))
    return docs


def clean_text(text: str) -> str:
    """Clean text by normalizing whitespace and line breaks.

    Args:
        text: Input text to clean

    Returns:
        Cleaned text string
    """
    if not text:
        return ""
    text = text.replace("\u3000", " ")
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def reflow_paragraphs(text: str) -> str:
    """Intelligently remove hard line breaks within paragraphs and merge incorrectly split paragraphs.

    The function:
    1) Splits by blank lines first; within a paragraph, if the previous line doesn't end with
       sentence-ending punctuation, merge it with the next line.
    2) If a paragraph doesn't end with sentence-ending punctuation and the next paragraph
       appears to be a continuation, merge across blank lines.
    3) Handles trailing hyphen word breaks.

    Args:
        text: Input text to reflow

    Returns:
        Reflowed text string
    """
    if not text:
        return ""

    text = text.replace("\r\n", "\n").replace("\r", "\n")

    end_punct_re = re.compile(r"[。！？!?；;…]\s*[”’」』》）】]*\s*$")
    next_start_re = re.compile(r'^[\u4e00-\u9fff0-9a-zA-Z“"‘’《（(【\[「『<]')

    def merge_lines_within_paragraph(para: str) -> str:
        lines = para.split("\n")
        segs: List[str] = []
        for ln in lines:
            ln = ln.strip()
            if not ln:
                continue
            if not segs:
                segs.append(ln)
                continue
            prev = segs[-1]
            should_join = not end_punct_re.search(prev)
            if should_join:
                if prev.endswith("-") and len(prev) > 1:
                    segs[-1] = prev[:-1] + ln
                else:
                    segs[-1] = prev + " " + ln
            else:
                segs.append(ln)
        joined = " ".join(segs)
        return re.sub(r"\s{2,}", " ", joined).strip()

    # First pass: merge lines within paragraphs
    raw_paras = re.split(r"\n{2,}", text)
    paras = [merge_lines_within_paragraph(p) for p in raw_paras if p.strip()]

    # Second pass: merge across paragraphs (handle incorrect blank lines causing sentence breaks)
    merged: List[str] = []
    for p in paras:
        if not merged:
            merged.append(p)
            continue
        prev = merged[-1]
        if prev and (not end_punct_re.search(prev)) and next_start_re.match(p):
            connector = "" if prev.endswith("-") else " "
            merged[-1] = re.sub(
                r"\s{2,}", " ", (prev.rstrip("-") + connector + p).strip()
            )
        else:
            merged.append(p)

    return "\n\n".join(merged).strip()


@app.tool(output="parse_file_path,text_corpus_save_path->None")
async def build_text_corpus(
    parse_file_path: str,
    text_corpus_save_path: str,
) -> None:
    """Build text corpus from various file formats.

    Args:
        parse_file_path: Path to file or directory containing files to parse
        text_corpus_save_path: Path where the text corpus JSONL file will be saved

    Raises:
        ToolError: If input path doesn't exist or required dependencies are missing
    """
    TEXT_EXTS = [".txt", ".md"]
    PMLIKE_EXT = [".pdf", ".xps", ".oxps", ".epub", ".mobi", ".fb2"]
    DOCX_EXT = [".docx"]
    WORD_LEGACY_EXT = [".doc", ".wps"]

    # Validate and sanitize path to prevent path traversal
    try:
        safe_path = _validate_path(parse_file_path)
        in_path = str(safe_path)
    except ValueError as e:
        err_msg = f"Invalid file path: {e}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    
    if not os.path.exists(in_path):
        err_msg = f"Input path not found: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    rows: List[Dict[str, Any]] = []
    is_single_input = os.path.isfile(in_path)

    def process_one_file(fp: str) -> None:
        fp_path = Path(fp)
        ext = fp_path.suffix.lower()
        stem = fp_path.stem
        content = ""

        if ext in TEXT_EXTS:
            try:
                from charset_normalizer import from_path
            except ImportError:
                err_msg = "charset_normalizer not installed. Please `pip install charset-normalizer`."
                app.logger.error(err_msg)
                raise ToolError(err_msg)
            try:
                results = from_path(fp).best()
                content = str(results) if results else ""
            except Exception as e:
                app.logger.warning(f"Text read failed: {fp} | {e}")

        elif ext in DOCX_EXT:
            try:
                content = _read_via_office_convert(fp)
                if content is None:
                    content = _read_docx_text(fp)
                if content is None:
                    err_msg = f"Unable to parse .docx file: {fp}. Unsupported file format."
                    if is_single_input:
                        app.logger.error(err_msg)
                        raise ToolError(err_msg)
                    app.logger.warning(err_msg)
                    content = ""
            except ToolError:
                raise
            except Exception as e:
                app.logger.warning(f"Docx read failed: {fp} | {e}")

        elif ext in WORD_LEGACY_EXT:
            content = _read_via_office_convert(fp)
            if content is None:
                err_msg = (
                    "Legacy Word format requires LibreOffice/soffice conversion. "
                    f"Please install LibreOffice or convert to .docx first: {fp}. "
                    "Unsupported file format."
                )
                if is_single_input:
                    app.logger.error(err_msg)
                    raise ToolError(err_msg)
                app.logger.warning(err_msg)
                return

        elif ext in PMLIKE_EXT:
            try:
                import pymupdf
            except ImportError:
                err_msg = "pymupdf not installed. Please `pip install pymupdf`."
                app.logger.error(err_msg)
                raise ToolError(err_msg)
            try:
                doc = None
                with suppress_stdout():
                    doc = pymupdf.open(fp)
                    texts = []
                    for pg in doc:
                        blocks = pg.get_text("blocks")
                        blocks.sort(key=lambda b: (b[1], b[0]))
                        page_text = "\n".join([b[4] for b in blocks if b[4].strip()])
                        texts.append(page_text)
                    content = "\n\n".join(texts)
            except Exception as e:
                app.logger.warning(f"PDF read failed: {fp} | {e}")
            finally:
                # Ensure PDF document is closed to prevent memory leaks
                if doc is not None:
                    try:
                        doc.close()
                    except Exception:
                        pass
        else:
            err_msg = f"Unsupported file type: {fp}. Unsupported file format."
            if is_single_input:
                app.logger.error(err_msg)
                raise ToolError(err_msg)
            app.logger.warning(err_msg)
            return

        if content.strip():
            rows.append(
                {
                    "id": stem,
                    "title": stem,
                    "contents": reflow_paragraphs(clean_text(content)),
                }
            )

    if os.path.isfile(in_path):
        process_one_file(in_path)
    else:
        all_files = []
        for dp, _, fns in os.walk(in_path):
            for fn in sorted(fns):
                all_files.append(os.path.join(dp, fn))

        for fp in tqdm(all_files, desc="Building text corpus", unit="file"):
            process_one_file(fp)

    out_path = os.path.abspath(text_corpus_save_path)
    _save_jsonl(rows, out_path)

    info_msg = (
        f"Built text corpus: {out_path} "
        f"(rows={len(rows)}, from={'dir' if os.path.isdir(in_path) else 'file'}: {in_path})"
    )
    app.logger.info(info_msg)


@app.tool(output="parse_file_path,image_corpus_save_path->None")
async def build_image_corpus(
    parse_file_path: str,
    image_corpus_save_path: str,
) -> None:
    """Build image corpus from PDF files by extracting pages as images.

    Args:
        parse_file_path: Path to PDF file or directory containing PDF files
        image_corpus_save_path: Path where the image corpus JSONL file will be saved

    Raises:
        ToolError: If input path doesn't exist, no PDFs found, or pymupdf is not installed
    """
    try:
        import pymupdf
    except ImportError:
        err_msg = "pymupdf not installed. Please `pip install pymupdf`."
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    # Validate and sanitize path to prevent path traversal
    try:
        safe_path = _validate_path(parse_file_path)
        in_path = str(safe_path)
    except ValueError as e:
        err_msg = f"Invalid file path: {e}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    
    if not os.path.exists(in_path):
        err_msg = f"Input path not found: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    # Validate output path
    try:
        safe_output_path = _validate_path(image_corpus_save_path)
        corpus_jsonl = str(safe_output_path)
    except ValueError as e:
        err_msg = f"Invalid output path: {e}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    out_root = os.path.dirname(corpus_jsonl) or os.getcwd()
    base_img_dir = os.path.join(out_root, "image")
    os.makedirs(base_img_dir, exist_ok=True)

    pdf_list: List[str] = []
    if os.path.isfile(in_path):
        if not in_path.lower().endswith(".pdf"):
            err_msg = f"Only PDF is supported here. Got: {os.path.splitext(in_path)[1]}"
            app.logger.error(err_msg)
            raise ToolError(err_msg)
        pdf_list = [in_path]
    else:
        for dp, _, fns in os.walk(in_path):
            for fn in sorted(fns):
                if fn.lower().endswith(".pdf"):
                    pdf_list.append(os.path.join(dp, fn))
        pdf_list.sort()

    if not pdf_list:
        err_msg = f"No PDF files found under: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    valid_rows: List[Dict[str, Any]] = []
    gid = 0

    for pdf_path in tqdm(pdf_list, desc="Building image corpus", unit="pdf"):
        stem = os.path.splitext(os.path.basename(pdf_path))[0]
        out_img_dir = os.path.join(base_img_dir, stem)
        os.makedirs(out_img_dir, exist_ok=True)

        doc = None
        try:
            with suppress_stdout():
                doc = pymupdf.open(pdf_path)
        except Exception as e:
            warn_msg = f"Skip PDF (open failed): {pdf_path} | reason: {e}"
            app.logger.warning(warn_msg)
            continue

        if doc is None:
            continue

        if getattr(doc, "is_encrypted", False):
            try:
                doc.authenticate("")
            except Exception:
                warn_msg = f"Skip PDF (encrypted): {pdf_path}"
                app.logger.warning(warn_msg)
                continue

        zoom = 144 / 72.0
        mat = pymupdf.Matrix(zoom, zoom)

        for i, page in enumerate(doc):
            try:
                with suppress_stdout():
                    pix = page.get_pixmap(
                        matrix=mat, alpha=False, colorspace=pymupdf.csRGB
                    )
            except Exception as e:
                warn_msg = f"Skip page {i} in {pdf_path}: render error: {e}"
                app.logger.warning(warn_msg)
                continue

            filename = f"page_{i}.jpg"
            save_path = os.path.join(out_img_dir, filename)
            rel_path = Path(os.path.join("image", stem, filename)).as_posix()

            try:
                pix.save(save_path, jpg_quality=90)
            except Exception as e:
                warn_msg = f"Skip page {i} in {pdf_path}: save error: {e}"
                app.logger.warning(warn_msg)
                continue
            finally:
                pix = None

            try:
                with Image.open(save_path) as im:
                    im.verify()
            except Exception as e:
                warn_msg = f"Skip page {i} in {pdf_path}: invalid image after save: {e}"
                app.logger.warning(warn_msg)
                try:
                    os.remove(save_path)
                except OSError as e:
                    warn_msg = f"Skip page {i} in {pdf_path}: remove error: {e}"
                    app.logger.warning(warn_msg)
                continue

            valid_rows.append(
                {
                    "id": gid,
                    "image_id": Path(os.path.join(stem, filename)).as_posix(),
                    "image_path": rel_path,
                }
            )
            gid += 1
        
        # Ensure PDF document is closed to prevent memory leaks
        if doc is not None:
            try:
                doc.close()
            except Exception:
                pass

    _save_jsonl(valid_rows, corpus_jsonl)
    info_msg = (
        f"Built image corpus: {corpus_jsonl} (valid images={len(valid_rows)}), "
        f"images root: {base_img_dir}, "
        f"pdf_count={len(pdf_list)}"
    )
    app.logger.info(info_msg)


@app.tool(output="parse_file_path,mineru_dir,mineru_extra_params->None")
async def mineru_parse(
    parse_file_path: str,
    mineru_dir: str,
    mineru_extra_params: Optional[Dict[str, Any]] = None,
) -> None:
    """Parse PDF files using MinerU tool.

    Args:
        parse_file_path: Path to PDF file or directory containing PDF files
        mineru_dir: Output directory for MinerU parsing results
        mineru_extra_params: Optional dictionary of extra parameters for MinerU command

    Raises:
        ToolError: If mineru executable not found, invalid input path, or execution fails
    """
    if shutil.which("mineru") is None:
        err_msg = "`mineru` executable not found. Please install it or add it to PATH."
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    if not parse_file_path:
        err_msg = "`parse_file_path` cannot be empty."
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    # Validate and sanitize path to prevent path traversal
    try:
        safe_path = _validate_path(parse_file_path)
        in_path = str(safe_path)
    except ValueError as e:
        err_msg = f"Invalid file path: {e}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    
    if not os.path.exists(in_path):
        err_msg = f"Input path not found: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    if os.path.isfile(in_path) and not in_path.lower().endswith(".pdf"):
        err_msg = f"Only .pdf files or directories are supported: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    out_root = os.path.abspath(mineru_dir)
    os.makedirs(out_root, exist_ok=True)

    proc_env = os.environ.copy()
    extra_args: List[str] = []
    if mineru_extra_params:
        source_val = mineru_extra_params.get("source")
        if source_val:
            # MinerU's internal model downloader reads MINERU_MODEL_SOURCE.
            # Passing --source alone may not affect download source in newer versions.
            proc_env["MINERU_MODEL_SOURCE"] = str(source_val)
            app.logger.info(
                "Set MINERU_MODEL_SOURCE=%s from mineru_extra_params.source",
                source_val,
            )
        for k in sorted(mineru_extra_params.keys()):
            if k == "source":
                continue
            v = mineru_extra_params[k]
            extra_args.append(f"--{k}")
            if v is not None and v != "":
                extra_args.append(str(v))

    cmd = ["mineru", "-p", in_path, "-o", out_root] + extra_args
    info_msg = f"Starting mineru command: {' '.join(cmd)}"
    app.logger.info(info_msg)

    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            env=proc_env,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.STDOUT,
        )
        assert proc.stdout is not None
        use_live_tqdm = _can_render_live_tqdm()
        progress_state: Dict[str, int] = {}
        progress_bars: Dict[str, tqdm] = {}
        last_plain_line = ""
        suppressed_progress = 0
        async for line in proc.stdout:
            for cleaned in _iter_clean_subprocess_lines(line):
                progress = _extract_progress_update(cleaned)
                if progress:
                    name, pct = progress
                    last_pct = progress_state.get(name, 0)
                    if use_live_tqdm:
                        bar = progress_bars.get(name)
                        if bar is None:
                            bar = tqdm(
                                total=100,
                                desc=f"[mineru] {name}",
                                unit="%",
                                dynamic_ncols=True,
                                leave=False,
                            )
                            progress_bars[name] = bar
                        if pct < last_pct:
                            bar.n = 0
                            bar.refresh()
                            last_pct = 0
                        delta = max(0, pct - last_pct)
                        if delta:
                            bar.update(delta)
                            progress_state[name] = pct
                        if pct >= 100:
                            bar.set_postfix_str("done")
                            bar.refresh()
                            bar.close()
                            progress_bars.pop(name, None)
                    else:
                        # Keep logs readable for non-interactive sessions.
                        if pct >= 100 or pct - last_pct >= 10:
                            progress_state[name] = pct
                            app.logger.info(f"[mineru] {name}: {pct}%")
                    suppressed_progress += 1
                    continue

                if cleaned == last_plain_line:
                    continue
                last_plain_line = cleaned
                app.logger.info(f"[mineru] {cleaned}")

        for bar in progress_bars.values():
            with suppress(Exception):
                bar.close()

        if suppressed_progress > 0:
            app.logger.info(
                f"[mineru] Suppressed {suppressed_progress} verbose progress lines"
            )

        returncode = await proc.wait()
        if returncode != 0:
            err_msg = f"mineru exited with non-zero code: {returncode}"
            app.logger.error(err_msg)
            raise ToolError(err_msg)
    except Exception as e:
        err_msg = f"Unexpected error while running mineru: {e}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    info_msg = f"mineru finished processing {in_path} into {out_root}"
    app.logger.info(info_msg)


def _list_images(images_dir: str) -> List[str]:
    """List all image files in a directory recursively.

    Args:
        images_dir: Directory path to search for images

    Returns:
        List of relative paths to image files (using forward slashes)
    """
    if not os.path.isdir(images_dir):
        return []
    exts = {".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff"}
    rels = []
    for dp, _, fns in os.walk(images_dir):
        for fn in sorted(fns):
            if os.path.splitext(fn)[1].lower() in exts:
                rel = os.path.relpath(os.path.join(dp, fn), start=images_dir)
                rels.append(Path(rel).as_posix())
    rels.sort()
    return rels


@app.tool(
    output="mineru_dir,parse_file_path,text_corpus_save_path,image_corpus_save_path->None"
)
async def build_mineru_corpus(
    mineru_dir: str,
    parse_file_path: str,
    text_corpus_save_path: str,
    image_corpus_save_path: str,
) -> None:
    """Build text and image corpus from MinerU parsing results.

    Args:
        mineru_dir: Directory containing MinerU parsing results
        parse_file_path: Original path to PDF file(s) that were parsed
        text_corpus_save_path: Path where the text corpus JSONL file will be saved
        image_corpus_save_path: Path where the image corpus JSONL file will be saved

    Raises:
        ToolError: If mineru_dir doesn't exist, invalid input path, or no PDFs found
    """
    root = os.path.abspath(mineru_dir)
    if not os.path.isdir(root):
        err_msg = f"MinerU root not found: {root}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    if not parse_file_path:
        err_msg = "`parse_file_path` cannot be empty."
        app.logger.error(err_msg)
        raise ToolError(err_msg)
    in_path = os.path.abspath(parse_file_path)
    if not os.path.exists(in_path):
        err_msg = f"Input path not found: {in_path}"
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    stems: List[str] = []
    if os.path.isfile(in_path):
        if not in_path.lower().endswith(".pdf"):
            err_msg = f"Only .pdf supported for file input: {in_path}"
            app.logger.error(err_msg)
            raise ToolError(err_msg)
        stems = [os.path.splitext(os.path.basename(in_path))[0]]
    else:
        seen: Set[str] = set()
        for dp, _, fns in os.walk(in_path):
            for fn in sorted(fns):
                if fn.lower().endswith(".pdf"):
                    stem = os.path.splitext(fn)[0]
                    if stem not in seen:
                        stems.append(stem)
                        seen.add(stem)
        stems.sort()
        if not stems:
            err_msg = f"No PDF files found under: {in_path}"
            app.logger.error(err_msg)
            raise ToolError(err_msg)

    text_rows: List[Dict[str, Any]] = []
    image_rows: List[Dict[str, Any]] = []
    image_out = os.path.abspath(image_corpus_save_path)
    out_root_dir = os.path.dirname(image_out)
    base_out_img_dir = os.path.join(out_root_dir, "images")
    os.makedirs(base_out_img_dir, exist_ok=True)

    for stem in stems:
        auto_dir = os.path.join(root, stem, "auto")
        if not os.path.isdir(auto_dir):
            warn_msg = f"Auto dir not found for '{stem}': {auto_dir} (skip)"
            app.logger.warning(warn_msg)
            continue

        md_path = os.path.join(auto_dir, f"{stem}.md")
        if not os.path.isfile(md_path):
            warn_msg = f"Markdown not found for '{stem}': {md_path} (skip text)"
            app.logger.warning(warn_msg)
        else:
            with open(md_path, "r", encoding="utf-8") as f:
                md_text = f.read().strip()
            text_rows.append({"id": stem, "title": stem, "contents": md_text})

        images_dir = os.path.join(auto_dir, "images")
        if not os.path.isdir(images_dir):
            warn_msg = f"No images dir for '{stem}': {images_dir} (skip images)"
            app.logger.warning(warn_msg)
            continue

        rel_list = _list_images(images_dir)
        for idx, rel in enumerate(rel_list):
            src = os.path.join(images_dir, rel)
            dst = os.path.join(base_out_img_dir, stem, rel)
            os.makedirs(os.path.dirname(dst), exist_ok=True)

            try:
                with Image.open(src) as im:
                    im.convert("RGB").copy()
            except Exception as e:
                warn_msg = f"Skip invalid image for '{stem}': {src}, reason: {e}"
                app.logger.warning(warn_msg)
                continue

            shutil.copy2(src, dst)
            image_rows.append(
                {
                    "id": len(image_rows),
                    "image_id": Path(os.path.join(stem, rel)).as_posix(),
                    "image_path": Path(os.path.join("images", stem, rel)).as_posix(),
                }
            )

    text_out = os.path.abspath(text_corpus_save_path)
    _save_jsonl(text_rows, text_out)
    _save_jsonl(image_rows, image_out)

    info_msg = (
        f"Built MinerU corpus from {in_path} | docs={len(stems)} | "
        f"text_rows={len(text_rows)} | image_rows={len(image_rows)}\n"
        f"Text corpus -> {text_out}\n"
        f"Image corpus -> {image_out} (images root: {base_out_img_dir})"
    )
    app.logger.info(info_msg)


@app.tool(
    output="raw_chunk_path,chunk_backend_configs,chunk_backend,tokenizer_or_token_counter,chunk_size,chunk_path,use_title->None"
)
async def chunk_documents(
    raw_chunk_path: str,
    chunk_backend_configs: Dict[str, Any],
    chunk_backend: str = "token",
    tokenizer_or_token_counter: str = "character",
    chunk_size: int = 256,
    chunk_path: Optional[str] = None,
    use_title: bool = True,
) -> None:
    """Chunk documents using various chunking strategies.

    Args:
        raw_chunk_path: Path to JSONL file containing documents to chunk
        chunk_backend_configs: Dictionary of configuration for each chunking backend
        chunk_backend: Chunking method to use ("token", "sentence", or "recursive")
        tokenizer_or_token_counter: Tokenizer name or counter type ("word", "character", or tiktoken encoding)
        chunk_size: Target size for each chunk
        chunk_path: Optional output path for chunked documents (defaults to project output directory)
        use_title: Whether to include document title in chunk content

    Raises:
        ToolError: If chonkie/tiktoken not installed, invalid chunking method, or chunking fails
    """
    try:
        from chonkie import (
            TokenChunker,
            SentenceChunker,
            RecursiveChunker,
            RecursiveRules,
        )
        import tiktoken
    except ImportError:
        err_msg = (
            "chonkie or tiktoken not installed. Please `pip install chonkie tiktoken`."
        )
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    if chunk_path is None:
        current_file = os.path.abspath(__file__)
        project_root = os.path.dirname(os.path.dirname(current_file))
        output_dir = os.path.join(project_root, "output", "corpus")
        chunk_path = os.path.join(output_dir, "chunks.jsonl")
    else:
        chunk_path = str(chunk_path)
        output_dir = os.path.dirname(chunk_path)
    os.makedirs(output_dir, exist_ok=True)
    documents = _load_jsonl(raw_chunk_path)

    cfg = (chunk_backend_configs.get(chunk_backend) or {}).copy()

    tokenizer_name = tokenizer_or_token_counter
    if tokenizer_name not in ["word", "character"]:
        try:
            tokenizer = tiktoken.get_encoding(tokenizer_name)
        except Exception:
            app.logger.warning(
                f"Could not load tokenizer '{tokenizer_name}', falling back to 'gpt2'."
            )
            tokenizer = tiktoken.get_encoding("gpt2")
    else:
        tokenizer = tokenizer_name

    chunk_overlap = cfg.get("chunk_overlap", 64)

    if chunk_overlap >= chunk_size:
        app.logger.warning(
            f"chunk_overlap ({chunk_overlap}) >= chunk_size ({chunk_size}), "
            "adjusting overlap to size/4."
        )
        chunk_overlap = int(chunk_size / 4)

    app.logger.info(
        f"Chunking Config: backend={chunk_backend}, size={chunk_size}, "
        f"overlap={chunk_overlap}, tokenizer={tokenizer_name}"
    )

    if chunk_backend == "token":
        chunker = TokenChunker(
            tokenizer=tokenizer,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
        )

    elif chunk_backend == "sentence":
        min_sentences_per_chunk = cfg.get("min_sentences_per_chunk", 1)

        delim = cfg.get("delim")
        DELIM_DEFAULT = [".", "!", "?", "；", "。", "！", "？"]
        if isinstance(delim, str):
            try:
                delim = ast.literal_eval(delim)
            except Exception:
                delim = DELIM_DEFAULT
        elif delim is None:
            delim = DELIM_DEFAULT

        chunker = SentenceChunker(
            tokenizer=tokenizer,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            min_sentences_per_chunk=min_sentences_per_chunk,
            delim=delim,
        )

    elif chunk_backend == "recursive":
        min_characters_per_chunk = cfg.get("min_characters_per_chunk", 50)

        chunker = RecursiveChunker(
            tokenizer=tokenizer,
            chunk_size=chunk_size,
            rules=RecursiveRules(),
            min_characters_per_chunk=min_characters_per_chunk,
        )

    else:
        err_msg = (
            f"Invalid chunking method: {chunk_backend}. "
            "Supported: token, sentence, recursive."
        )
        app.logger.error(err_msg)
        raise ToolError(err_msg)

    chunked_documents = []
    current_chunk_id = 0
    for doc in tqdm(documents, desc=f"Chunking ({chunk_backend})", unit="doc"):
        doc_id = doc.get("id") or ""
        title = (doc.get("title") or "").strip()
        text = (doc.get("contents") or "").strip()

        if not text:
            warn_msg = f"doc_id={doc_id} has no contents, skipped."
            app.logger.warning(warn_msg)
            continue
        try:
            chunks = chunker.chunk(text)
        except Exception as e:
            err_msg = f"fail chunked(doc_id={doc_id}): {e}"
            app.logger.error(err_msg)
            raise ToolError(err_msg)

        for chunk in chunks:
            if use_title:
                contents = f"Title:\n{title}\n\nContent:\n{chunk.text}"
            else:
                contents = chunk.text
            meta_chunk = {
                "id": current_chunk_id,
                "doc_id": doc_id,
                "title": title,
                "contents": contents.strip(),
            }
            chunked_documents.append(meta_chunk)
            current_chunk_id += 1

    _save_jsonl(chunked_documents, chunk_path)


if __name__ == "__main__":
    app.run(transport="stdio")
