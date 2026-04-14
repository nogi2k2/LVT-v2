"""
vv_plan_parser.py
==================
Parse a V&V Plan .docx document and extract:
  - Document metadata (title, ER number, requirement doc, sample size)
  - Label PRD entries  (symbols: MD, UDI, Date of Manufacture, IFU)
  - IFU PRD entries    (IFU requirement text and section references)

Returns a VVPlanData dataclass.

Strategy:
  1. Extract all text from the docx (paragraphs + table cells).
  2. Parse the Verification Span table first to get PRD IDs, requirement
     text and V&V method directly from the structured table rows.
  3. Identify additional PRD IDs using the standard regex pattern.
  4. Classify each PRD as "label" or "IFU" using keyword patterns.
  5. Parse metadata from the header section of the document.
"""

import os
import re
import logging
from dataclasses import dataclass, field
from typing import List, Optional

logger = logging.getLogger(__name__)

# ── PRD ID pattern ────────────────────────────────────────────────────────────
PRD_ID_RE = re.compile(r"\b([A-Z]+(?:_[A-Z0-9]+-?[A-Z0-9]*)?_PRD\d+)\b")

# ── V&V Method keywords ───────────────────────────────────────────────────────
VV_METHOD_RE = re.compile(
    r"\b(Inspection|Analysis|Test(?:ing)?|Design|Demonstration|Similarity|Calculation)\b",
    re.IGNORECASE,
)

# ── Symbol keyword patterns ───────────────────────────────────────────────────
SYMBOL_PATTERNS = {
    "MD": [
        r"MD\s+symbol", r"medical\s+device\s+symbol",
        r"5\s*[.,]\s*7\s*[.,]\s*7",
        r"\bMD\b.*symbol", r"symbol.*\bMD\b",
    ],
    "UDI": [
        r"\bUDI\b", r"unique\s+device\s+identifier",
        r"5\s*[.,]\s*7\s*[.,]\s*10",
    ],
    "Date of Manufacture": [
        r"date\s+of\s+manufacture", r"date\s+of\s+mfg", r"manufacturing\s+date",
        r"5\s*[.,]\s*7\s*[.,]\s*1",
    ],
    "IFU": [
        r"consult\s+instructions?\s+for\s+use", r"\bIFU\b",
        r"5\s*[.,]\s*4\s*[.,]\s*3",
        r"instructions?\s+for\s+use",
    ],
}

# ── IFU section / requirement patterns ───────────────────────────────────────
IFU_SECTION_PATTERNS = [
    r"IFU",
    r"instructions?\s+for\s+use",
    r"addendum",
    r"ifu.*verif",
    r"verif.*ifu",
    r"section\s+\d+.*ifu",
]

# ── Metadata extraction patterns ──────────────────────────────────────────────
ER_RE      = re.compile(r"\bER\s*[-:]?\s*(\d{5,})", re.IGNORECASE)
SAMPLE_RE  = re.compile(r"sample\s+size\s*[:\-=]?\s*(\d+)", re.IGNORECASE)
REQ_DOC_RE = re.compile(r"requirement\s+doc(?:ument)?\s*[:\-]?\s*([\w\s\-\.]+)", re.IGNORECASE)


# ── Data classes ──────────────────────────────────────────────────────────────

@dataclass
class PRDEntry:
    prd_id:            str
    requirement_text:  str
    section:           str = ""
    symbols:           List[str] = field(default_factory=list)
    entry_type:        str = "label"    # "label" or "ifu"
    verification_type: str = "label"    # user-facing alias of entry_type
    vv_method:         str = ""         # e.g. "Inspection", "Analysis"
    raw_text:          str = ""


@dataclass
class VVPlanData:
    doc_title:   str = ""
    er_number:   str = ""
    req_doc:     str = ""
    sample_size: str = ""
    prd_entries: List[PRDEntry] = field(default_factory=list)
    label_prds:  List[PRDEntry] = field(default_factory=list)
    ifu_prds:    List[PRDEntry] = field(default_factory=list)
    raw_error:   str = ""


# ── Helpers ───────────────────────────────────────────────────────────────────

def _matches_any(text: str, patterns: list) -> bool:
    for p in patterns:
        try:
            if re.search(p, text, re.IGNORECASE):
                return True
        except re.error:
            if p.lower() in text.lower():
                return True
    return False


def _detect_symbols(text: str) -> List[str]:
    """Return list of symbol names matched in text."""
    found = []
    for sym, patterns in SYMBOL_PATTERNS.items():
        if _matches_any(text, patterns):
            found.append(sym)
    return found


def _is_ifu_entry(text: str, section: str = "") -> bool:
    """Return True if this PRD entry relates to IFU verification."""
    combined = f"{section} {text}"
    return _matches_any(combined, IFU_SECTION_PATTERNS)


def _extract_vv_method(text: str) -> str:
    """Extract the V&V method keyword from a text fragment."""
    m = VV_METHOD_RE.search(text)
    if m:
        return m.group(1).capitalize()
    return ""


# ── Table-based Verification Span extraction ──────────────────────────────────

def _extract_from_tables(doc) -> dict:
    """
    Walk all tables in the docx looking for the Verification Span table.

    The expected structure (any column ordering that contains a PRD ID):
        S.No | Requirement ID     | Requirement Text | V&V Method
        1    | ACC_MPV_PRD147     | The Circuit ...  | Inspection

    Returns a dict: prd_id -> {"requirement_text": str, "vv_method": str}
    """
    results = {}
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # Detect header row: find which column contains "requirement" / "PRD"
        header_cells = [c.text.strip().lower() for c in rows[0].cells]

        # Find column indices by header keyword
        prd_col  = next((i for i, h in enumerate(header_cells)
                         if "requirement id" in h or "prd" in h or "req" in h), None)
        text_col = next((i for i, h in enumerate(header_cells)
                         if "requirement text" in h or "text" in h
                         or "description" in h), None)
        meth_col = next((i for i, h in enumerate(header_cells)
                         if "v&v" in h or "method" in h or "verification" in h), None)

        for row in rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue

            # Find a PRD ID anywhere in the row if column detection failed
            row_text = " ".join(cells)
            prd_matches = PRD_ID_RE.findall(row_text)
            if not prd_matches:
                continue

            for prd_id in prd_matches:
                req_text = ""
                if text_col is not None and text_col < len(cells):
                    req_text = cells[text_col]
                else:
                    # Use any cell that isn't just the PRD ID
                    for cell in cells:
                        if prd_id not in cell and len(cell) > 10:
                            req_text = cell
                            break

                vv_method = ""
                if meth_col is not None and meth_col < len(cells):
                    vv_method = _extract_vv_method(cells[meth_col]) or cells[meth_col].strip()
                else:
                    vv_method = _extract_vv_method(row_text)

                if prd_id not in results:
                    results[prd_id] = {
                        "requirement_text": req_text[:600],
                        "vv_method":        vv_method,
                    }

    return results


# ── Paragraph/text block extraction ──────────────────────────────────────────

def _extract_all_text(docx_path: str):
    """
    Return (doc, list_of_(section_hint, text)) from the docx.
    section_hint is the nearest heading above this paragraph.
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required to parse V&V Plans.\n"
            "Install with:  pip install python-docx"
        )

    doc = Document(docx_path)
    blocks = []
    current_section = ""

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            current_section = text
        blocks.append((current_section, text))

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                combined = " | ".join(row_texts)
                blocks.append((current_section, combined))

    return doc, blocks


def _extract_metadata(blocks: list) -> dict:
    meta = {"doc_title": "", "er_number": "", "req_doc": "", "sample_size": ""}

    for i, (section, text) in enumerate(blocks[:40]):
        if not meta["doc_title"] and len(text) > 8 and len(text) < 150:
            if any(kw in text.lower() for kw in ["v&v", "verification", "plan", "protocol"]):
                meta["doc_title"] = text

        if not meta["er_number"]:
            m = ER_RE.search(text)
            if m:
                meta["er_number"] = f"ER{m.group(1)}"

        if not meta["sample_size"]:
            m = SAMPLE_RE.search(text)
            if m:
                meta["sample_size"] = m.group(1)

        if not meta["req_doc"]:
            m = REQ_DOC_RE.search(text)
            if m:
                meta["req_doc"] = m.group(1).strip()[:60]

    if not meta["doc_title"] and blocks:
        meta["doc_title"] = blocks[0][1][:80]

    return meta


def _extract_prd_entries(blocks: list, table_data: dict) -> List[PRDEntry]:
    """
    Walk all blocks, identify PRD IDs, build PRDEntry for each.
    Merges in structured data from table_data (Verification Span table).
    """
    entries: dict = {}

    all_text = "\n".join(t for _, t in blocks)
    id_positions = [(m.start(), m.group(1)) for m in PRD_ID_RE.finditer(all_text)]

    for i, (pos, prd_id) in enumerate(id_positions):
        end = id_positions[i+1][0] if i+1 < len(id_positions) else pos + 800
        raw_chunk = all_text[pos:end].replace("\n", " ").strip()
        raw_chunk = re.sub(r"\s{2,}", " ", raw_chunk)

        req_text = raw_chunk[len(prd_id):].strip()
        req_text = re.sub(r"^(Requirement|Heading|Information)\s*", "", req_text,
                          flags=re.IGNORECASE)
        req_text = req_text[:400]

        if prd_id in entries:
            continue

        # Prefer structured table data if available
        td = table_data.get(prd_id, {})
        if td.get("requirement_text"):
            req_text = td["requirement_text"]
        vv_method = td.get("vv_method", "") or _extract_vv_method(req_text)

        section_hint = ""
        for sec, txt in blocks:
            if prd_id in txt:
                section_hint = sec
                break

        syms       = _detect_symbols(req_text)
        entry_type = "ifu" if _is_ifu_entry(req_text, section_hint) else "label"

        entries[prd_id] = PRDEntry(
            prd_id            = prd_id,
            requirement_text  = req_text,
            section           = section_hint,
            symbols           = syms,
            entry_type        = entry_type,
            verification_type = entry_type,
            vv_method         = vv_method,
            raw_text          = raw_chunk[:600],
        )

    return list(entries.values())


# ── Public API ────────────────────────────────────────────────────────────────

def parse_vv_plan(docx_path: str) -> VVPlanData:
    """
    Parse a V&V Plan .docx and return a VVPlanData object.
    Never raises — errors are captured in VVPlanData.raw_error.
    """
    data = VVPlanData()

    if not os.path.isfile(docx_path):
        data.raw_error = f"File not found: {docx_path}"
        return data

    try:
        doc, blocks = _extract_all_text(docx_path)
        table_data  = _extract_from_tables(doc)
        meta        = _extract_metadata(blocks)
        entries     = _extract_prd_entries(blocks, table_data)

        data.doc_title   = meta["doc_title"]
        data.er_number   = meta["er_number"]
        data.req_doc     = meta["req_doc"]
        data.sample_size = meta["sample_size"]
        data.prd_entries = entries
        data.label_prds  = [e for e in entries if e.entry_type == "label"]
        data.ifu_prds    = [e for e in entries if e.entry_type == "ifu"]

        logger.info(
            "V&V Plan parsed: %s  |  %d PRDs  |  %d label  |  %d IFU",
            os.path.basename(docx_path),
            len(entries), len(data.label_prds), len(data.ifu_prds),
        )

    except Exception as e:
        logger.exception("V&V Plan parse failed: %s", e)
        data.raw_error = str(e)

    return data
